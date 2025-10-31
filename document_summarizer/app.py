"""
Advanced Document Summarizer with Dark/Light Theme and Enhanced Text Cleaning
"""

import streamlit as st
import re
import numpy as np
from collections import Counter
import io
from datetime import datetime
import nltk
import requests
import urllib.parse
from config import Config

# Download required NLTK data for cloud deployment
@st.cache_resource
def download_nltk_data():
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt', quiet=True)
    
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords', quiet=True)
    
    try:
        nltk.data.find('corpora/wordnet')
    except LookupError:
        nltk.download('wordnet', quiet=True)

# Download NLTK data on startup
download_nltk_data()

# File processing imports with better error handling
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError as e:
    st.error(f"PyPDF2 not available: {e}")
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError as e:
    st.error(f"python-docx not available: {e}")
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError as e:
    REPORTLAB_AVAILABLE = False

class AdvancedTextCleaner:
    """Advanced text cleaning to remove headers, footers, and unwanted content."""
    
    def __init__(self):
        # Enhanced patterns for headers and footers
        self.header_patterns = [
            r'KESHAV MEMORIAL INSTITU?TE OF TECHNOLOGY.*?INSTITUTION\)',
            r'WEB TECHNOLOGIES.*?PAGE NO:\d+',
            r'UNIT\s*-\s*[IVX]+.*?REGULATIONS',
            r'KR\d+\s+Regulations.*?INSTITUTION\)',
            r'KMIT.*?AUTONOMOUS INSTITUTION.*?PAGE NO:\d+',
            r'AN AUTONOMO?US INSTITUTION.*?(?=\n)',
            r'PAGE NO:\d+.*?(?=\n)',
            r'Syllabus:.*?(?=\n\n|\n[A-Z])',
            r'Figure:.*?(?=\n)',
            r'Example.*?:.*?(?=\n\n)',
            r'Filename:.*?(?=\n)',
            r'Output.*?:.*?(?=\n\n)',
            r'Browser.*?output.*?:.*?(?=\n\n)',
            r'Console.*?output.*?:.*?(?=\n\n)',
            r'Terminal.*?output.*?:.*?(?=\n\n)',
        ]
        
        # Enhanced repetitive patterns
        self.repetitive_patterns = [
            r'KESHAV MEMORIAL.*?(?=\n)',
            r'WEB TECHNOLOGIES.*?(?=\n)',
            r'UNIT\s*-\s*[IVX]+.*?(?=\n)',
            r'KR\d+.*?(?=\n)',
            r'KMIT.*?(?=\n)',
            r'AN AUTONOMO?US.*?(?=\n)',
            r'INSTITUTION.*?(?=\n)',
            r'PAGE NO:\d+.*?(?=\n)',
            r'Run:.*?js.*?(?=\n)',
            r'node.*?js.*?(?=\n)',
            r'npm.*?(?=\n)',
            r'Browser.*?url:.*?(?=\n)',
            r'localhost:\d+.*?(?=\n)',
            r'Question:.*?(?=\n)',
            r'Query:.*?(?=\n)',
            r'Write a query.*?(?=\n)',
            r'Q-?\d+\..*?(?=\n)',
        ]
        
        # Patterns for code blocks and technical jargon
        self.code_patterns = [
            r'const\s+\w+.*?;',
            r'require\s*\(.*?\)',
            r'console\.log\s*\(.*?\)',
            r'app\.\w+\s*\(.*?\)',
            r'res\.\w+\s*\(.*?\)',
            r'req\.\w+.*?(?=\s)',
            r'db\.\w+.*?(?=\s)',
            r'mongod.*?(?=\s)',
            r'mongosh.*?(?=\s)',
        ]
    
    def clean_academic_content(self, text):
        """Remove academic headers, footers, and repetitive content."""
        if not isinstance(text, str):
            return ""
        
        # Remove header patterns (more aggressive)
        for pattern in self.header_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE | re.DOTALL)
        
        # Remove repetitive patterns (more aggressive)
        for pattern in self.repetitive_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)
        
        # Remove specific institutional references
        institutional_patterns = [
            r'KESHAV MEMORIAL.*?TECHNOLOGY',
            r'AN AUTONOMOUS INSTITUTION',
            r'INSTITU?TE OF TECHNOLOGY',
            r'\(AN AUTONOMO?US INSTITUTION\)',
        ]
        
        for pattern in institutional_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        
        # Clean up excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple newlines to double
        text = re.sub(r'\s+', ' ', text)  # Multiple spaces to single
        
        # Remove standalone numbers and page references
        text = re.sub(r'\b\d+\s*\)', '', text)  # Remove numbered lists like "1)"
        text = re.sub(r'\bpage\s+\d+\b', '', text, flags=re.IGNORECASE)
        text = re.sub(r'PAGE NO:\d+', '', text, flags=re.IGNORECASE)
        
        # Remove URLs and email patterns
        text = re.sub(r'http[s]?://\S+', '', text)
        text = re.sub(r'\S+@\S+', '', text)
        
        # Remove excessive punctuation and special characters
        text = re.sub(r'[^\w\s\.\!\?\,\;\:\-\(\)]', ' ', text)
        
        # Remove very short fragments
        sentences = text.split('.')
        meaningful_sentences = [s.strip() for s in sentences if len(s.strip().split()) >= 3]
        text = '. '.join(meaningful_sentences)
        
        # Clean up and return
        return text.strip()
    
    def extract_meaningful_content(self, text):
        """Extract only meaningful sentences and paragraphs."""
        sentences = re.split(r'[.!?]+', text)
        meaningful_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            
            # Skip very short sentences
            if len(sentence.split()) < 5:
                continue
            
            # Skip sentences that are mostly technical jargon
            if self._is_mostly_technical(sentence):
                continue
            
            # Skip sentences with too many special characters
            if len(re.findall(r'[^\w\s]', sentence)) > len(sentence.split()) * 0.3:
                continue
            
            meaningful_sentences.append(sentence)
        
        return '. '.join(meaningful_sentences)
    
    def _is_mostly_technical(self, sentence):
        """Check if sentence is mostly technical code or jargon."""
        technical_indicators = [
            'const', 'require', 'console', 'app.', 'res.', 'req.',
            'function', 'var', 'let', 'npm', 'node', 'js',
            'localhost', 'port', 'server', 'http'
        ]
        
        words = sentence.lower().split()
        technical_count = sum(1 for word in words if any(tech in word for tech in technical_indicators))
        
        return technical_count > len(words) * 0.4

class URLProcessor:
    """Process documents from URLs and shared links."""
    
    def __init__(self):
        self.text_cleaner = AdvancedTextCleaner()
    
    def convert_google_drive_url(self, url):
        """Convert Google Drive sharing URL to direct download URL."""
        if "drive.google.com" in url:
            if "/file/d/" in url:
                # Extract file ID from sharing URL
                file_id = url.split("/file/d/")[1].split("/")[0]
                return f"https://drive.google.com/uc?export=download&id={file_id}"
        return url
    
    def convert_google_docs_url(self, url):
        """Convert Google Docs URL to plain text export URL."""
        if "docs.google.com" in url:
            if "/document/d/" in url:
                # Extract document ID
                doc_id = url.split("/document/d/")[1].split("/")[0]
                return f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
        return url
    
    def download_from_url(self, url):
        """Download content from URL."""
        try:
            # Convert special URLs
            if "drive.google.com" in url:
                url = self.convert_google_drive_url(url)
            elif "docs.google.com" in url:
                url = self.convert_google_docs_url(url)
            
            # Download with headers to mimic browser
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            return response.content, response.headers.get('content-type', '')
            
        except Exception as e:
            return None, f"Error downloading from URL: {str(e)}"
    
    def process_url_content(self, content, content_type, url):
        """Process downloaded content based on type."""
        try:
            # Determine file type from content-type or URL
            if 'pdf' in content_type.lower() or url.lower().endswith('.pdf'):
                return self._process_pdf_content(content)
            elif 'text' in content_type.lower() or url.lower().endswith('.txt'):
                return self._process_text_content(content)
            elif 'document' in content_type.lower() or 'docs.google.com' in url:
                return self._process_text_content(content)
            else:
                # Try to process as text
                return self._process_text_content(content)
                
        except Exception as e:
            return None, f"Error processing content: {str(e)}"
    
    def _process_pdf_content(self, content):
        """Process PDF content from bytes."""
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            # Apply cleaning
            cleaned_text = self.text_cleaner.clean_academic_content(text)
            meaningful_text = self.text_cleaner.extract_meaningful_content(cleaned_text)
            
            return meaningful_text.strip(), None
            
        except Exception as e:
            return None, f"Error processing PDF: {str(e)}"
    
    def _process_text_content(self, content):
        """Process text content from bytes."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                return None, "Could not decode text content"
            
            # Apply cleaning
            cleaned_text = self.text_cleaner.clean_academic_content(text)
            meaningful_text = self.text_cleaner.extract_meaningful_content(cleaned_text)
            
            return meaningful_text.strip(), None
            
        except Exception as e:
            return None, f"Error processing text: {str(e)}"

class DocumentProcessor:
    """Process different document formats with advanced cleaning."""
    
    def __init__(self):
        self.text_cleaner = AdvancedTextCleaner()
        self.url_processor = URLProcessor()
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract and clean text from PDF file."""
        if not PDF_AVAILABLE:
            return None, "PyPDF2 not installed. Install with: pip install PyPDF2"
        
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            # Apply advanced cleaning
            cleaned_text = self.text_cleaner.clean_academic_content(text)
            meaningful_text = self.text_cleaner.extract_meaningful_content(cleaned_text)
            
            return meaningful_text.strip(), None
        except Exception as e:
            return None, f"Error reading PDF: {str(e)}"
    
    def extract_text_from_docx(self, docx_file):
        """Extract and clean text from DOCX file."""
        if not DOCX_AVAILABLE:
            return None, "python-docx not installed. Install with: pip install python-docx"
        
        try:
            doc = Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Apply advanced cleaning
            cleaned_text = self.text_cleaner.clean_academic_content(text)
            meaningful_text = self.text_cleaner.extract_meaningful_content(cleaned_text)
            
            return meaningful_text.strip(), None
        except Exception as e:
            return None, f"Error reading DOCX: {str(e)}"
    
    def extract_text_from_txt(self, txt_file):
        """Extract and clean text from TXT file."""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    txt_file.seek(0)
                    text = txt_file.read().decode(encoding)
                    
                    # Apply advanced cleaning
                    cleaned_text = self.text_cleaner.clean_academic_content(text)
                    meaningful_text = self.text_cleaner.extract_meaningful_content(cleaned_text)
                    
                    return meaningful_text.strip(), None
                except UnicodeDecodeError:
                    continue
            return None, "Could not decode text file. Please ensure it's in UTF-8 format."
        except Exception as e:
            return None, f"Error reading TXT: {str(e)}"
    
    def process_uploaded_file(self, uploaded_file):
        """Process uploaded file and extract clean text."""
        if uploaded_file is None:
            return None, "No file uploaded"
        
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return self.extract_text_from_pdf(uploaded_file)
        elif file_extension == 'docx':
            return self.extract_text_from_docx(uploaded_file)
        elif file_extension == 'txt':
            return self.extract_text_from_txt(uploaded_file)
        else:
            return None, f"Unsupported file format: {file_extension}"
    
    def process_document_url(self, url):
        """Process document from URL."""
        if not url or not url.strip():
            return None, "No URL provided"
        
        url = url.strip()
        
        # Validate URL
        if not (url.startswith('http://') or url.startswith('https://')):
            return None, "Invalid URL. Please provide a complete URL starting with http:// or https://"
        
        # Download content
        content, error_or_content_type = self.url_processor.download_from_url(url)
        
        if content is None:
            return None, error_or_content_type
        
        # Process content
        return self.url_processor.process_url_content(content, error_or_content_type, url)

class SummaryFormatter:
    """Format summaries with proper structure and headings."""
    
    def format_summary(self, summary_text, original_filename):
        """Format summary in clean, readable paragraph style."""
        # Clean the summary text first
        summary_text = self._clean_summary_text(summary_text)
        
        # Split into sentences more intelligently
        sentences = self._split_into_sentences(summary_text)
        
        if not sentences:
            return f"## Summary of {original_filename}\n\nNo meaningful content could be extracted."
        
        if len(sentences) <= 2:
            return f"## Summary of {original_filename}\n\n{' '.join(sentences)}"
        
        # Create clean, well-structured summary
        formatted_summary = f"## Summary of {original_filename}\n\n"
        
        # Group sentences into coherent paragraphs
        paragraphs = self._create_paragraphs(sentences)
        
        # Join paragraphs with proper spacing
        formatted_summary += '\n\n'.join(paragraphs)
        
        return formatted_summary.strip()
    
    def _clean_summary_text(self, text):
        """Clean and normalize summary text."""
        import re
        
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Remove redundant phrases
        redundant_phrases = [
            r'in summary,?\s*',
            r'to summarize,?\s*',
            r'in conclusion,?\s*',
            r'overall,?\s*',
            r'this document\s*',
            r'the text\s*',
            r'the document\s*'
        ]
        
        for phrase in redundant_phrases:
            text = re.sub(phrase, '', text, flags=re.IGNORECASE)
        
        # Fix capitalization after cleaning
        if text and not text[0].isupper():
            text = text[0].upper() + text[1:]
        
        return text
    
    def _split_into_sentences(self, text):
        """Split text into meaningful sentences."""
        import re
        
        # Split on sentence boundaries
        sentences = re.split(r'[.!?]+', text)
        
        # Clean and filter sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            
            # Skip very short or meaningless sentences
            if len(sentence.split()) < 4:
                continue
            
            # Ensure proper capitalization
            if sentence and not sentence[0].isupper():
                sentence = sentence[0].upper() + sentence[1:]
            
            cleaned_sentences.append(sentence)
        
        return cleaned_sentences
    
    def _create_paragraphs(self, sentences):
        """Group sentences into coherent paragraphs."""
        if len(sentences) <= 3:
            # Single paragraph for short summaries
            return ['. '.join(sentences) + '.']
        
        elif len(sentences) <= 6:
            # Two paragraphs for medium summaries
            mid_point = len(sentences) // 2
            para1 = '. '.join(sentences[:mid_point]) + '.'
            para2 = '. '.join(sentences[mid_point:]) + '.'
            return [para1, para2]
        
        else:
            # Three paragraphs for longer summaries
            third = len(sentences) // 3
            para1 = '. '.join(sentences[:third]) + '.'
            para2 = '. '.join(sentences[third:third*2]) + '.'
            para3 = '. '.join(sentences[third*2:]) + '.'
            return [para1, para2, para3]

class SmartTextRankSummarizer:
    """Enhanced TextRank implementation with better sentence selection."""
    
    def __init__(self):
        self.formatter = SummaryFormatter()
    
    def segment_sentences(self, text):
        """Split text into meaningful sentences."""
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip().split()) >= 8]
        return sentences
    
    def _tokenize(self, text):
        """Advanced tokenization with stopword filtering."""
        import re
        tokens = re.findall(r'\w+', text.lower())
        
        # Basic stopwords
        stopwords = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
            'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
            'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
            'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those'
        }
        
        return [token for token in tokens if token not in stopwords and len(token) > 2]
    
    def _compute_tf_idf(self, sentences):
        """Compute TF-IDF vectors for sentences."""
        tokenized_sentences = [self._tokenize(sent) for sent in sentences]
        
        # Build vocabulary
        vocab = set()
        for tokens in tokenized_sentences:
            vocab.update(tokens)
        vocab = list(vocab)
        vocab_to_idx = {word: i for i, word in enumerate(vocab)}
        
        # Compute TF-IDF
        tfidf_vectors = []
        
        # Document frequency
        df = Counter()
        for tokens in tokenized_sentences:
            unique_tokens = set(tokens)
            for token in unique_tokens:
                df[token] += 1
        
        for tokens in tokenized_sentences:
            # Term frequency
            tf = Counter(tokens)
            total_terms = len(tokens)
            
            # TF-IDF vector
            vector = np.zeros(len(vocab))
            for token, count in tf.items():
                if token in vocab_to_idx:
                    tf_score = count / total_terms
                    idf_score = np.log(len(sentences) / (df[token] + 1))
                    vector[vocab_to_idx[token]] = tf_score * idf_score
            
            tfidf_vectors.append(vector)
        
        return np.array(tfidf_vectors)
    
    def _cosine_similarity(self, vec1, vec2):
        """Compute cosine similarity between two vectors."""
        dot_product = np.dot(vec1, vec2)
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)
        
        if norm1 == 0 or norm2 == 0:
            return 0
        
        return dot_product / (norm1 * norm2)
    
    def _build_similarity_matrix(self, sentences):
        """Build similarity matrix between sentences."""
        tfidf_vectors = self._compute_tf_idf(sentences)
        n = len(sentences)
        similarity_matrix = np.zeros((n, n))
        
        for i in range(n):
            for j in range(n):
                if i != j:
                    similarity_matrix[i][j] = self._cosine_similarity(
                        tfidf_vectors[i], tfidf_vectors[j]
                    )
        
        return similarity_matrix
    
    def _textrank_scores(self, similarity_matrix, damping=0.85, max_iter=100, tol=1e-4):
        """Compute TextRank scores."""
        n = similarity_matrix.shape[0]
        
        # Normalize similarity matrix
        row_sums = similarity_matrix.sum(axis=1)
        normalized_matrix = np.zeros_like(similarity_matrix)
        
        for i in range(n):
            if row_sums[i] > 0:
                normalized_matrix[i] = similarity_matrix[i] / row_sums[i]
        
        # Initialize scores
        scores = np.ones(n) / n
        
        # Iterate
        for _ in range(max_iter):
            new_scores = (1 - damping) / n + damping * normalized_matrix.T.dot(scores)
            
            if np.abs(new_scores - scores).sum() < tol:
                break
            
            scores = new_scores
        
        return scores
    
    def summarize(self, text, compression_ratio=0.3, filename="document"):
        """Generate smart extractive summary using TextRank with improved coherence."""
        sentences = self.segment_sentences(text)
        
        if len(sentences) <= 3:
            raw_summary = ' '.join(sentences)
        else:
            # Calculate number of sentences for coherent summary
            base_sentences = max(6, int(len(sentences) * compression_ratio))
            
            # Ensure we get enough sentences for coherent paragraphs
            if compression_ratio <= 0.2:
                num_sentences = max(6, base_sentences)  # At least 6 for coherence
            elif compression_ratio <= 0.4:
                num_sentences = max(8, base_sentences)  # At least 8 for balanced
            else:
                num_sentences = max(10, base_sentences)  # At least 10 for detailed
            
            num_sentences = min(num_sentences, len(sentences))
            
            try:
                # Build similarity matrix and compute scores
                similarity_matrix = self._build_similarity_matrix(sentences)
                scores = self._textrank_scores(similarity_matrix)
                
                # Get top sentences with better distribution for coherence
                top_indices = np.argsort(scores)[-num_sentences:][::-1]
                
                # Ensure we include first sentence for context
                if 0 not in top_indices and len(sentences) > 3:
                    top_indices = np.append(top_indices[:-1], 0)
                
                # Include last sentence if document is long enough
                if (len(sentences)-1) not in top_indices and len(sentences) > 8:
                    top_indices = np.append(top_indices[:-1], len(sentences)-1)
                
                # Sort to maintain document flow
                top_indices = sorted(top_indices)
                
                # Select sentences and create coherent flow
                summary_sentences = [sentences[i] for i in top_indices]
                
                # Improve sentence connections for better flow
                coherent_sentences = self._improve_sentence_flow(summary_sentences)
                raw_summary = '. '.join(coherent_sentences) + '.'
                
            except Exception as e:
                st.error(f"TextRank failed: {e}")
                # Better fallback with document structure awareness
                step = max(1, len(sentences) // num_sentences)
                selected_indices = list(range(0, len(sentences), step))[:num_sentences]
                raw_summary = '. '.join([sentences[i] for i in selected_indices]) + '.'
        
        # Format the summary with proper structure
        formatted_summary = self.formatter.format_summary(raw_summary, filename)
        return formatted_summary
    
    def _improve_sentence_flow(self, sentences):
        """Improve sentence flow and coherence."""
        if len(sentences) <= 2:
            return sentences
        
        # Remove redundant information and improve transitions
        improved_sentences = []
        prev_words = set()
        
        for sentence in sentences:
            words = set(sentence.lower().split())
            # Check for significant overlap with previous sentences
            overlap = len(words.intersection(prev_words)) / len(words) if words else 0
            
            if overlap < 0.7:  # Keep sentences with less than 70% overlap
                improved_sentences.append(sentence)
                prev_words.update(words)
        
        return improved_sentences if improved_sentences else sentences

@st.cache_resource
def load_components():
    """Load components with caching."""
    try:
        doc_processor = DocumentProcessor()
        summarizer = SmartTextRankSummarizer()
        return doc_processor, summarizer
    except Exception as e:
        st.error(f"Error loading components: {e}")
        # Fallback without caching
        return DocumentProcessor(), SmartTextRankSummarizer()

def load_transformer_models():
    """Load transformer models with caching and API key support."""
    try:
        # Try to import transformers (optional dependency)
        from transformers import pipeline
        import warnings
        warnings.filterwarnings("ignore", category=UserWarning)
        
        # Check for Hugging Face API key
        hf_token = Config.get_hf_token()
        
        # Create summarization pipelines with working models
        models = {}
        
        # List of working models (in order of preference)
        model_configs = [
            ('t5', 't5-small'),
        ]
        
        # Try to load additional models if T5 works
        additional_models = [
            ('bart', 'facebook/bart-large-cnn'),
            ('flan-t5', 'google/flan-t5-small')
        ]
        
        loaded_models = []
        
        # Load T5 first (most reliable) - try without token first
        for model_name, model_path in model_configs:
            try:
                # Try without token first (public models)
                model = pipeline("summarization", model=model_path)
                models[model_name] = model
                loaded_models.append(model_name.upper())
                
            except Exception as e:
                # If that fails, try with token if available
                if hf_token:
                    try:
                        model = pipeline("summarization", model=model_path, token=hf_token)
                        models[model_name] = model
                        loaded_models.append(model_name.upper())
                    except Exception as e2:
                        continue
                else:
                    continue
        
        # Try to load additional models (optional)
        for model_name, model_path in additional_models:
            try:
                # Try without token first (public models)
                model = pipeline("summarization", model=model_path)
                models[model_name] = model
                loaded_models.append(model_name.upper())
                
            except Exception as e:
                # If that fails, try with token if available
                if hf_token:
                    try:
                        model = pipeline("summarization", model=model_path, token=hf_token)
                        models[model_name] = model
                        loaded_models.append(model_name.upper())
                    except Exception as e2:
                        continue
                else:
                    continue
        
        if models:
            st.success(f"🤖 Ready with {len(models)} model(s): {', '.join(loaded_models)}")
            return models
        else:
            st.info("💡 Transformers available but models couldn't load - using TextRank!")
            return None
        
    except ImportError:
        st.info("💡 Transformers not installed - using enhanced TextRank algorithm!")
        return None
    except Exception as e:
        st.info("💡 Using enhanced TextRank as fallback.")
        return None
    return DocumentProcessor(), SmartTextRankSummarizer()

def apply_custom_css():
    """Apply custom CSS for better UI."""
    st.markdown("""
    <style>
    /* Main container styling */
    .main > div {
        padding-top: 2rem;
    }
    
    /* Style the file uploader container */
    .stFileUploader {
        border-radius: 20px;
        padding: 3rem 2rem;
        text-align: center;
        margin: 2rem 0;
        transition: all 0.3s ease;
        cursor: pointer;
    }
    
    .stFileUploader:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.15);
    }
    
    /* File uploader text styling */
    .stFileUploader label {
        font-size: 1.1rem;
        font-weight: 500;
    }
    
    /* Button styling */
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(0,0,0,0.3);
    }
    
    /* Metrics styling */
    .metric-container {
        background: rgba(255,255,255,0.1);
        padding: 1rem;
        border-radius: 10px;
        backdrop-filter: blur(10px);
        border: 1px solid rgba(255,255,255,0.2);
    }
    
    /* Summary box styling */
    .summary-box {
        background: linear-gradient(135deg, rgba(255,255,255,0.1) 0%, rgba(255,255,255,0.05) 100%);
        padding: 2rem;
        border-radius: 15px;
        border: 1px solid rgba(255,255,255,0.2);
        backdrop-filter: blur(10px);
        margin: 1rem 0;
    }
    
    /* Improved button styling */
    .stButton > button {
        transition: all 0.3s ease;
        font-weight: 600;
    }
    
    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    
    /* Header styling */
    .main-header {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(135deg, rgba(255,255,255,0.1) 0%, rgba(255,255,255,0.05) 100%);
        border-radius: 20px;
        margin-bottom: 2rem;
        backdrop-filter: blur(10px);
        border: 1px solid rgba(255,255,255,0.2);
    }
    
    /* Upload area styling */
    .upload-area {
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
        border: 2px dashed rgba(102, 126, 234, 0.5);
        border-radius: 20px;
        padding: 3rem 2rem;
        text-align: center;
        margin: 2rem 0;
        transition: all 0.3s ease;
    }
    
    .upload-area:hover {
        border-color: rgba(102, 126, 234, 0.8);
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.15) 0%, rgba(118, 75, 162, 0.15) 100%);
    }
    </style>
    """, unsafe_allow_html=True)

def create_download_buttons(text, base_filename):
    """Create download buttons for different formats."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # TXT download
        txt_filename = f"{base_filename}_summary_{timestamp}.txt"
        txt_buffer = io.BytesIO()
        txt_buffer.write(text.encode('utf-8'))
        txt_buffer.seek(0)
        
        st.download_button(
            label="📄 Download TXT",
            data=txt_buffer.getvalue(),
            file_name=txt_filename,
            mime="text/plain",
            type="primary",
            use_container_width=True
        )
    
    with col2:
        # PDF download (using reportlab if available, otherwise HTML to PDF)
        if REPORTLAB_AVAILABLE:
            try:
                pdf_filename = f"{base_filename}_summary_{timestamp}.pdf"
                pdf_buffer = io.BytesIO()
                
                doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Add title
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=16,
                    spaceAfter=30,
                )
                story.append(Paragraph(f"Summary of {base_filename}", title_style))
                story.append(Spacer(1, 12))
                
                # Add content
                for line in text.split('\n'):
                    if line.strip():
                        if line.startswith('#'):
                            # Header
                            story.append(Paragraph(line.replace('#', '').strip(), styles['Heading2']))
                        elif line.startswith('•') or line.startswith('-'):
                            # Bullet point
                            story.append(Paragraph(line, styles['Normal']))
                        else:
                            # Regular text
                            story.append(Paragraph(line, styles['Normal']))
                        story.append(Spacer(1, 6))
                
                doc.build(story)
                pdf_buffer.seek(0)
                
                st.download_button(
                    label="📕 Download PDF",
                    data=pdf_buffer.getvalue(),
                    file_name=pdf_filename,
                    mime="application/pdf",
                    type="primary",
                    use_container_width=True
                )
            except Exception as e:
                st.button("📕 PDF (Error)", disabled=True, use_container_width=True, help=f"PDF generation error: {e}")
        else:
            st.button("📕 PDF (Not Available)", disabled=True, use_container_width=True, help="ReportLab not installed")
    
    with col3:
        # DOCX download
        try:
            from docx import Document
            from docx.shared import Inches
            
            docx_filename = f"{base_filename}_summary_{timestamp}.docx"
            docx_buffer = io.BytesIO()
            
            doc = Document()
            doc.add_heading(f'Summary of {base_filename}', 0)
            
            for line in text.split('\n'):
                if line.strip():
                    if line.startswith('# '):
                        doc.add_heading(line.replace('#', '').strip(), level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line.replace('##', '').strip(), level=2)
                    elif line.startswith('•') or line.startswith('-'):
                        p = doc.add_paragraph(line.replace('•', '').replace('-', '').strip())
                        p.style = 'List Bullet'
                    elif line.strip() and not line.startswith('#'):
                        doc.add_paragraph(line.strip())
            
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            st.download_button(
                label="📘 Download DOCX",
                data=docx_buffer.getvalue(),
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
        except ImportError:
            st.button("📘 DOCX (Available)", disabled=True, use_container_width=True)

def main():
    st.set_page_config(
        page_title="Smart Document Summarizer",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="collapsed"
    )
    
    # Hide sidebar completely
    st.markdown("""
    <style>
    .css-1d391kg {display: none}
    .css-1rs6os {display: none}
    .css-17eq0hr {display: none}
    section[data-testid="stSidebar"] {display: none}
    </style>
    """, unsafe_allow_html=True)
    
    # Apply custom CSS
    apply_custom_css()
    

    

    
    # Apply clean theme styles
    st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
        color: white;
    }
    .stFileUploader {
        background: linear-gradient(135deg, rgba(59, 130, 246, 0.1) 0%, rgba(147, 51, 234, 0.1) 100%) !important;
        border: 2px dashed rgba(59, 130, 246, 0.5) !important;
    }
    .stFileUploader:hover {
        border-color: rgba(59, 130, 246, 0.8) !important;
        background: linear-gradient(135deg, rgba(59, 130, 246, 0.15) 0%, rgba(147, 51, 234, 0.15) 100%) !important;
    }
    .main-header {
        background: linear-gradient(135deg, rgba(30, 41, 59, 0.8) 0%, rgba(51, 65, 85, 0.8) 100%) !important;
        border: 1px solid rgba(59, 130, 246, 0.2) !important;
    }
    .summary-box {
        background: linear-gradient(135deg, rgba(30, 41, 59, 0.6) 0%, rgba(51, 65, 85, 0.6) 100%) !important;
        border: 1px solid rgba(59, 130, 246, 0.3) !important;
    }
    .metric-container {
        background: rgba(30, 41, 59, 0.6) !important;
        border: 1px solid rgba(59, 130, 246, 0.2) !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Main header
    st.markdown("""
    <div class="main-header">
        <h1>🤖 Smart Document Summarizer</h1>
        <p>Upload files or paste URLs to get AI-powered summaries with advanced text cleaning</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Load components (try cached first, fallback to direct)
    try:
        doc_processor, summarizer = load_components()
    except Exception as e:
        st.warning(f"Using fallback components due to: {e}")
        doc_processor = DocumentProcessor()
        summarizer = SmartTextRankSummarizer()
    
    # Create two columns for layout
    col1, col2 = st.columns([1, 1], gap="large")
    
    with col1:
        st.markdown("### 📁 Upload Your Document")
        
        uploaded_file = st.file_uploader(
            "📄 Drag & Drop Your File Here or Click to Browse\n\nSupports PDF, DOCX, and TXT files",
            type=['pdf', 'docx', 'txt'],
            help="Supported formats: PDF, DOCX, TXT"
        )
        
        st.markdown("**OR**")
        
        # URL input for shared documents
        document_url = st.text_input(
            "🔗 Enter Document URL",
            placeholder="https://drive.google.com/file/d/... or https://docs.google.com/document/d/...",
            help="Paste Google Drive, Google Docs, Dropbox, or direct file links"
        )
        
        # Summarization method selection
        st.markdown("### 🤖 AI Model Selection")
        summarization_method = st.selectbox(
            "Choose Summarization Approach",
            ["🧠 T5 Transformer (Default)", "📚 BART Transformer", "🚀 FLAN-T5 (Advanced)", "🔧 TextRank (Fallback)"],
            help="Select the AI model for summarization"
        )
        
        # Compression ratio slider
        st.markdown("### ⚙️ Summary Settings")
        compression_ratio = st.slider(
            "Summary Length",
            min_value=0.1,
            max_value=0.8,
            value=0.3,
            step=0.1,
            help="Lower values = shorter summaries"
        )
        
        # Display method info
        if "TextRank" in summarization_method:
            st.info("⚡ Fast extractive summarization using graph-based ranking")
        elif "T5" in summarization_method:
            st.info("🧠 Advanced abstractive summarization using T5 transformer (Default)")
        elif "BART" in summarization_method:
            st.info("📚 High-quality abstractive summarization using BART transformer")
        elif "FLAN-T5" in summarization_method:
            st.info("🚀 Advanced instruction-tuned T5 model for better summarization")
        else:
            st.info("🤖 AI-powered summarization with transformer models")
    
    with col2:
        st.markdown("### 📊 Document Analysis")
        
        # Handle both file upload and URL input
        extracted_text = None
        error = None
        original_name = "document"
        
        if uploaded_file is not None:
            # Display file info with nice styling
            st.markdown(f"""
            <div class="metric-container">
                <h4>📄 {uploaded_file.name}</h4>
                <p>Size: {uploaded_file.size:,} bytes</p>
                <p>Type: {uploaded_file.type}</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Process the uploaded file
            with st.spinner("🔍 Extracting and cleaning text..."):
                extracted_text, error = doc_processor.process_uploaded_file(uploaded_file)
            original_name = uploaded_file.name.rsplit('.', 1)[0]
            
        elif document_url:
            # Display URL info
            st.markdown(f"""
            <div class="metric-container">
                <h4>🔗 Document URL</h4>
                <p>Source: {document_url[:50]}...</p>
                <p>Type: Web Document</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Process the URL
            with st.spinner("🌐 Downloading and processing document from URL..."):
                extracted_text, error = doc_processor.process_document_url(document_url)
            
            # Extract name from URL
            try:
                from urllib.parse import urlparse
                parsed_url = urlparse(document_url)
                original_name = parsed_url.path.split('/')[-1] or "web_document"
                if '.' in original_name:
                    original_name = original_name.rsplit('.', 1)[0]
            except:
                original_name = "web_document"
        
        # Handle processing results
        if error:
            st.error(f"❌ {error}")
            
            if "not installed" in error:
                st.markdown("""
                **📦 Install required packages:**
                ```bash
                pip install PyPDF2 python-docx requests
                ```
                """)
        
        elif extracted_text:
            # Show extracted text info
                # Show extracted text info
                word_count = len(extracted_text.split())
                char_count = len(extracted_text)
                
                # Metrics in a nice layout
                metric_col1, metric_col2 = st.columns(2)
                with metric_col1:
                    st.metric("📊 Words", f"{word_count:,}")
                with metric_col2:
                    st.metric("📝 Characters", f"{char_count:,}")
                
                # Show preview of cleaned text
                with st.expander("👀 Preview Cleaned Text"):
                    preview_text = extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text
                    st.text_area("Clean Document Content", preview_text, height=200)
                
                # Generate summary button
                if st.button("🚀 Generate Smart Summary", type="primary", use_container_width=True):
                    
                    # Load transformer models if needed
                    if "Transformer" in summarization_method or "FLAN-T5" in summarization_method:
                        with st.spinner("🤖 Loading AI models..."):
                            transformer_models = load_transformer_models()
                    else:
                        transformer_models = None
                    
                    with st.spinner("🧠 Creating intelligent summary with AI..."):
                        try:
                            if "T5" in summarization_method:
                                # Use T5 transformer (abstractive) - DEFAULT
                                available_models = ['t5', 'flan-t5']
                                model_found = None
                                
                                if transformer_models:
                                    for model_name in available_models:
                                        if model_name in transformer_models:
                                            model_found = transformer_models[model_name]
                                            break
                                
                                if model_found:
                                    max_length = max(100, int(len(extracted_text.split()) * compression_ratio * 2))
                                    min_length = max(30, max_length // 4)
                                    
                                    try:
                                        # Chunk text if too long
                                        if len(extracted_text.split()) > 1000:
                                            chunks = [extracted_text[i:i+3000] for i in range(0, len(extracted_text), 3000)]
                                            summaries = []
                                            for chunk in chunks[:3]:  # Process max 3 chunks
                                                chunk_summary = model_found(chunk, max_length=max_length//len(chunks), min_length=min_length//len(chunks))
                                                summaries.append(chunk_summary[0]['summary_text'])
                                            raw_summary = ' '.join(summaries)
                                        else:
                                            result = model_found(extracted_text, max_length=max_length, min_length=min_length)
                                            raw_summary = result[0]['summary_text']
                                        
                                        summary = summarizer.formatter.format_summary(raw_summary, original_name)
                                        st.info("🧠 Used T5 Transformer for advanced abstractive summarization (Default)")
                                    except Exception as e:
                                        st.warning(f"🔄 T5 generation failed: {str(e)[:50]}... using TextRank")
                                        summary = summarizer.summarize(extracted_text, compression_ratio, original_name)
                                else:
                                    # Fallback to TextRank
                                    st.warning("🔄 T5 model not available, using TextRank as fallback")
                                    summary = summarizer.summarize(extracted_text, compression_ratio, original_name)
                                
                            elif "BART" in summarization_method:
                                # Use BART transformer (abstractive)
                                available_models = ['bart']  # Only use working BART model
                                model_found = None
                                
                                if transformer_models:
                                    for model_name in available_models:
                                        if model_name in transformer_models:
                                            model_found = transformer_models[model_name]
                                            break
                                
                                if model_found:
                                    max_length = max(100, int(len(extracted_text.split()) * compression_ratio * 2))
                                    min_length = max(30, max_length // 4)
                                    
                                    try:
                                        # Chunk text if too long
                                        if len(extracted_text.split()) > 1000:
                                            chunks = [extracted_text[i:i+3000] for i in range(0, len(extracted_text), 3000)]
                                            summaries = []
                                            for chunk in chunks[:3]:
                                                chunk_summary = model_found(chunk, max_length=max_length//len(chunks), min_length=min_length//len(chunks))
                                                summaries.append(chunk_summary[0]['summary_text'])
                                            raw_summary = ' '.join(summaries)
                                        else:
                                            result = model_found(extracted_text, max_length=max_length, min_length=min_length)
                                            raw_summary = result[0]['summary_text']
                                        
                                        summary = summarizer.formatter.format_summary(raw_summary, original_name)
                                        st.info("📚 Used BART Transformer for high-quality abstractive summarization")
                                    except Exception as e:
                                        st.warning(f"🔄 BART generation failed: {str(e)[:50]}... using TextRank")
                                        summary = summarizer.summarize(extracted_text, compression_ratio, original_name)
                                else:
                                    # Fallback to TextRank
                                    st.warning("🔄 BART model not available, using TextRank as fallback")
                                    summary = summarizer.summarize(extracted_text, compression_ratio, original_name)
                                
                            elif "FLAN-T5" in summarization_method:
                                # Use FLAN-T5 transformer (advanced)
                                available_models = ['flan-t5']
                                model_found = None
                                
                                if transformer_models:
                                    for model_name in available_models:
                                        if model_name in transformer_models:
                                            model_found = transformer_models[model_name]
                                            break
                                
                                if model_found:
                                    max_length = max(100, int(len(extracted_text.split()) * compression_ratio * 2))
                                    min_length = max(30, max_length // 4)
                                    
                                    try:
                                        # Chunk text if too long
                                        if len(extracted_text.split()) > 1000:
                                            chunks = [extracted_text[i:i+3000] for i in range(0, len(extracted_text), 3000)]
                                            summaries = []
                                            for chunk in chunks[:3]:
                                                chunk_summary = model_found(chunk, max_length=max_length//len(chunks), min_length=min_length//len(chunks))
                                                summaries.append(chunk_summary[0]['summary_text'])
                                            raw_summary = ' '.join(summaries)
                                        else:
                                            result = model_found(extracted_text, max_length=max_length, min_length=min_length)
                                            raw_summary = result[0]['summary_text']
                                        
                                        summary = summarizer.formatter.format_summary(raw_summary, original_name)
                                        st.info("🚀 Used FLAN-T5 Transformer for advanced instruction-tuned summarization")
                                    except Exception as e:
                                        st.warning(f"🔄 FLAN-T5 generation failed: {str(e)[:50]}... using TextRank")
                                        summary = summarizer.summarize(extracted_text, compression_ratio, original_name)
                                else:
                                    # Fallback to TextRank
                                    st.warning("🔄 FLAN-T5 model not available, using TextRank as fallback")
                                    summary = summarizer.summarize(extracted_text, compression_ratio, original_name)
                                

                            elif "TextRank" in summarization_method:
                                # Use TextRank (extractive) - Fallback option
                                summary = summarizer.summarize(extracted_text, compression_ratio, original_name)
                                st.info("⚡ Used TextRank algorithm for fast extractive summarization")
                                
                            else:
                                # Default to any available transformer, otherwise TextRank
                                if transformer_models:
                                    # Use the first available model
                                    model_name = list(transformer_models.keys())[0]
                                    model_pipeline = transformer_models[model_name]
                                    
                                    max_length = max(100, int(len(extracted_text.split()) * compression_ratio * 2))
                                    min_length = max(30, max_length // 4)
                                    
                                    try:
                                        result = model_pipeline(extracted_text, max_length=max_length, min_length=min_length)
                                        raw_summary = result[0]['summary_text']
                                        summary = summarizer.formatter.format_summary(raw_summary, original_name)
                                        st.info(f"🧠 Used {model_name.upper()} Transformer (Default)")
                                    except Exception as e:
                                        st.warning(f"🔄 {model_name.upper()} failed: {str(e)[:30]}... using TextRank")
                                        summary = summarizer.summarize(extracted_text, compression_ratio, original_name)
                                else:
                                    summary = summarizer.summarize(extracted_text, compression_ratio, original_name)
                                    st.info("⚡ Used TextRank algorithm (Fallback)")
                                
                        except Exception as e:
                            st.error(f"❌ Error during summarization: {e}")
                            st.info("🔄 Falling back to TextRank algorithm")
                            summary = summarizer.summarize(extracted_text, compression_ratio, original_name)
                            summary = None
                    
                    if summary:
                        st.success("✅ Summary generated successfully!")
                        
                        # Display formatted summary in scrollable container
                        st.markdown("### 📋 Smart Summary")
                        
                        # Create scrollable container with improved CSS
                        st.markdown("""
                        <style>
                        .summary-container {
                            max-height: 500px;
                            overflow-y: auto;
                            padding: 30px;
                            border: 1px solid rgba(59, 130, 246, 0.3);
                            border-radius: 15px;
                            background: linear-gradient(135deg, rgba(30, 41, 59, 0.8) 0%, rgba(51, 65, 85, 0.8) 100%);
                            margin: 15px 0;
                            scrollbar-width: thin;
                            scrollbar-color: rgba(59, 130, 246, 0.5) transparent;
                            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.3);
                        }
                        
                        .summary-container::-webkit-scrollbar {
                            width: 10px;
                        }
                        
                        .summary-container::-webkit-scrollbar-track {
                            background: rgba(30, 41, 59, 0.3);
                            border-radius: 5px;
                        }
                        
                        .summary-container::-webkit-scrollbar-thumb {
                            background: linear-gradient(135deg, rgba(59, 130, 246, 0.6), rgba(147, 51, 234, 0.6));
                            border-radius: 5px;
                        }
                        
                        .summary-container::-webkit-scrollbar-thumb:hover {
                            background: linear-gradient(135deg, rgba(59, 130, 246, 0.8), rgba(147, 51, 234, 0.8));
                        }
                        
                        .summary-container p {
                            margin: 15px 0;
                            text-align: justify;
                            line-height: 1.8;
                        }
                        
                        .summary-container h3 {
                            border-bottom: 2px solid rgba(59, 130, 246, 0.3);
                            padding-bottom: 10px;
                            margin-bottom: 20px;
                        }
                        </style>
                        """, unsafe_allow_html=True)
                        
                        # Display summary in clean, readable format
                        # Parse the summary to extract title and content
                        summary_lines = summary.split('\n')
                        title = summary_lines[0].replace('##', '').strip() if summary_lines else "Summary"
                        content = '\n'.join(summary_lines[2:]) if len(summary_lines) > 2 else summary
                        
                        # Format content into proper paragraphs
                        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                        formatted_content = ""
                        
                        for paragraph in paragraphs:
                            formatted_content += f'<p style="margin: 15px 0; line-height: 1.8; text-align: justify;">{paragraph}</p>'
                        
                        # Display with clean formatting
                        st.markdown(f"""
                        <div class="summary-container">
                            <h3 style="color: #3b82f6; margin-bottom: 20px; font-size: 1.4em;">{title}</h3>
                            <div style="font-size: 1.1em;">
                                {formatted_content}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Summary statistics
                        summary_words = len(summary.split())
                        actual_compression = summary_words / word_count if word_count > 0 else 0
                        
                        stat_col1, stat_col2, stat_col3 = st.columns(3)
                        with stat_col1:
                            st.metric("📊 Summary Words", f"{summary_words:,}")
                        with stat_col2:
                            st.metric("📉 Compression", f"{actual_compression:.2f}")
                        with stat_col3:
                            st.metric("⚡ Reduction", f"{(1-actual_compression)*100:.1f}%")
                        
                        # Download buttons for multiple formats
                        st.markdown("### 📥 Download Summary")
                        create_download_buttons(summary, original_name)
                        
                    else:
                        st.error("❌ Failed to generate summary. Please try with a different document.")
        
        else:
            # Show instructions when no file is uploaded and no URL provided
            st.markdown("""
            <div class="summary-box">
                <h3>🚀 How to Use</h3>
                <p><strong>1.</strong> Upload your document (PDF, DOCX, TXT) or paste a URL</p>
                <p><strong>2.</strong> Choose your AI model (T5, BART, FLAN-T5)</p>
                <p><strong>3.</strong> Adjust the summary length slider</p>
                <p><strong>4.</strong> Click "Generate Smart Summary"</p>
                <p><strong>5.</strong> Download your clean, intelligent summary</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Feature highlights
            st.markdown("### ✨ Features")
            st.markdown("""
            - 🧹 **Advanced Text Cleaning**: Removes headers, footers, and academic jargon
            - 🤖 **AI Models**: T5, BART, FLAN-T5 transformers with TextRank fallback
            - 🔗 **URL Support**: Google Drive, Google Docs, Dropbox, and direct links
            - 🎨 **Beautiful UI**: Modern dark theme with professional styling
            - 📥 **Multi-Format Export**: Download as TXT, PDF, or DOCX
            - 🔒 **Privacy First**: All processing happens locally (except URL downloads)
            """)
    

    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; padding: 2rem; background: rgba(255,255,255,0.1); border-radius: 15px; backdrop-filter: blur(10px);'>
        <p>🤖 <strong>Smart Document Summarizer</strong> | Powered by Advanced TextRank Algorithm</p>
        <p>🔒 Your documents are processed locally - no data is stored or shared</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()